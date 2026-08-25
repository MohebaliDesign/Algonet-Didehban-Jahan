from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "research" / "report-source.md"
OUTPUT = ROOT / "Didehban-Jahan-Onboarding-Research.docx"

IMG_MAP = Path(r"C:\Users\p\AppData\Local\Temp\codex-clipboard-9e4e73e9-b77e-4ef9-ab43-ba95db108e87.png")
IMG_SOURCES = Path(r"C:\Users\p\AppData\Local\Temp\codex-clipboard-72e3336a-9d09-4f42-aa10-c2fc211630d8.png")

FONT = "Tahoma"
NAVY = "0B2545"
TEAL = "0B8F80"
MUTED = "566573"
LIGHT = "F2F4F7"
CALLOUT = "EAF7F4"
GRID = "D7DEE8"
WHITE = "FFFFFF"


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rtl = rpr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        rpr.append(rtl)
    rtl.set(qn("w:val"), "1")


def set_paragraph_rtl(paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        ppr.append(bidi)
    bidi.set(qn("w:val"), "1")
    paragraph.alignment = alignment


def set_keep_with_next(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    if ppr.find(qn("w:keepNext")) is None:
        ppr.append(OxmlElement("w:keepNext"))


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.find(qn("w:tcMar"))
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size="4"):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent_dxa))
    tblind.set(qn("w:type"), "dxa")
    bidi_visual = tblpr.find(qn("w:bidiVisual"))
    if bidi_visual is None:
        bidi_visual = OxmlElement("w:bidiVisual")
        tblpr.append(bidi_visual)
    bidi_visual.set(qn("w:val"), "1")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            width = widths_dxa[min(i, len(widths_dxa) - 1)]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def add_hyperlink(paragraph, text, url, color=TEAL):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT)
    rpr.append(rfonts)
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rpr.append(rtl)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rpr.append(c)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


TOKEN_RE = re.compile(r"(\*\*.*?\*\*|\[[^\]]+\]\([^)]+\))")


def add_inline_markup(paragraph, text, size=11, color=None, bold=False):
    pos = 0
    for match in TOKEN_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, color=color, bold=bold)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            lm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            add_hyperlink(paragraph, lm.group(1), lm.group(2))
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color, bold=bold)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, TEAL, 16, 8),
        ("Heading 2", 13, TEAL, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            style._element.rPr.rFonts.set(qn(f"w:{attr}"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for base_name, new_name in (("List Bullet", "RTL Bullet"), ("List Number", "RTL Number")):
        if new_name not in doc.styles:
            style = doc.styles.add_style(new_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles[base_name]
        else:
            style = doc.styles[new_name]
        style.font.name = FONT
        style.font.size = Pt(11)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            style._element.rPr.rFonts.set(qn(f"w:{attr}"), FONT)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    if "Image Caption RTL" not in doc.styles:
        cap = doc.styles.add_style("Image Caption RTL", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = doc.styles["Image Caption RTL"]
    cap.font.name = FONT
    cap.font.size = Pt(9)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor.from_string(MUTED)
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(10)


def add_header_footer(doc):
    section = doc.sections[0]
    hp = section.header.paragraphs[0]
    set_paragraph_rtl(hp, WD_ALIGN_PARAGRAPH.RIGHT)
    hr = hp.add_run("دیده‌بان جهان | گزارش انبوردینگ محصول")
    set_run_font(hr, size=8.5, color=MUTED)

    fp = section.footer.paragraphs[0]
    set_paragraph_rtl(fp, WD_ALIGN_PARAGRAPH.CENTER)
    r = fp.add_run("صفحه ")
    set_run_font(r, size=9, color=MUTED)
    add_page_field(fp)


def add_cover(doc):
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
    p = doc.add_paragraph()
    set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("گزارش پژوهش و انبوردینگ محصول")
    set_run_font(r, size=12, bold=True, color=TEAL)
    p.paragraph_format.space_after = Pt(14)

    p = doc.add_paragraph()
    set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("دیده‌بان جهان")
    set_run_font(r, size=30, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("تحلیل نسخه V15، معماری اطلاعات و سه بنچمارک اصلی")
    set_run_font(r, size=15, color=MUTED)
    p.paragraph_format.space_after = Pt(30)

    p = doc.add_paragraph()
    set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("از داده خام تا گزارش قابل اتکا، با منشأ و عدم قطعیت قابل مشاهده")
    set_run_font(r, size=12, italic=True, color=TEAL)
    p.paragraph_format.space_after = Pt(80)

    for text in ("برای: تیم محصول، طراحی و فنی دیده‌بان جهان", "تاریخ بررسی: ۲ شهریور ۱۴۰۵ / 24 August 2026", "وضعیت: مبنای تصمیم‌گیری برای بازطراحی و ساخت مرحله بعد"):
        p = doc.add_paragraph()
        set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(text)
        set_run_font(r, size=10, color=MUTED)
        p.paragraph_format.space_after = Pt(5)
    doc.add_page_break()


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="B7DDD5", size="6")
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.RIGHT)
    add_inline_markup(p, text, size=11, color=NAVY, bold=True)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_image(doc, path, caption, alt_text):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.25))
    for docpr in run._r.xpath(".//wp:docPr"):
        docpr.set("descr", alt_text)
    cap = doc.add_paragraph(style="Image Caption RTL")
    set_paragraph_rtl(cap, WD_ALIGN_PARAGRAPH.CENTER)
    add_inline_markup(cap, caption, size=9, color=MUTED)


def clean_inline_url(line):
    match = re.search(r"(https?://\S+)$", line)
    if match and "](http" not in line:
        url = match.group(1)
        return line[:match.start()] + f"[مشاهده منبع]({url})"
    return line


def widths_for_columns(count):
    presets = {
        2: [2700, 6660],
        3: [1900, 2600, 4860],
        4: [1500, 2200, 2500, 3160],
        5: [1400, 1300, 1900, 1700, 3060],
    }
    return presets.get(count, [9360 // count] * (count - 1) + [9360 - (9360 // count) * (count - 1)])


def add_markdown_table(doc, rows):
    if not rows:
        return
    count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=count)
    set_table_geometry(table, widths_for_columns(count))
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for ri, row in enumerate(rows):
        for ci in range(count):
            cell = table.cell(ri, ci)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.RIGHT)
            text = row[ci].strip() if ci < len(row) else ""
            add_inline_markup(p, text, size=9.2 if count >= 4 else 10, bold=(ri == 0), color=NAVY if ri == 0 else None)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


PAGE_BREAK_HEADINGS = {
    "تحلیل بنچمارک‌ها",
    "معماری اطلاعات پیشنهادی",
    "دفتر ادعا و منبع",
}


def parse_body(doc, markdown):
    lines = markdown.splitlines()
    start = next((i for i, line in enumerate(lines) if line.strip() == "## فرض‌ها"), 0)
    i = start
    while i < len(lines):
        raw = lines[i].rstrip()
        line = raw.strip()
        if not line:
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[i + 1].strip()):
            rows = []
            header = [c.strip() for c in line.strip("|").split("|")]
            rows.append(header)
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            add_markdown_table(doc, rows)
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            set_paragraph_rtl(p)
            add_inline_markup(p, line[4:], size=13, color=TEAL, bold=True)
            set_keep_with_next(p)
            i += 1
            continue
        if line.startswith("## "):
            title = line[3:]
            if title in PAGE_BREAK_HEADINGS and len(doc.paragraphs) > 5:
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            set_paragraph_rtl(p)
            add_inline_markup(p, title, size=16, color=TEAL, bold=True)
            set_keep_with_next(p)
            if title == "تحلیل نسخه V15 از روی تصاویر":
                add_image(doc, IMG_MAP, "نمای نقشه و لایه‌های مستقل در نسخه V15؛ قدرت پوشش بالا در کنار تراکم کنترل‌ها.", "اسکرین‌شات نسخه V15 دیده‌بان جهان با نقشه جهان و لایه‌های رخداد")
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue

        if re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph(style="RTL Number")
            set_paragraph_rtl(p)
            ppr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:right"), "720")
            ind.set(qn("w:hanging"), "360")
            ppr.append(ind)
            add_inline_markup(p, clean_inline_url(text))
            i += 1
            continue
        if line.startswith("- "):
            text = line[2:]
            p = doc.add_paragraph(style="RTL Bullet")
            set_paragraph_rtl(p)
            ppr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:right"), "720")
            ind.set(qn("w:hanging"), "360")
            ppr.append(ind)
            add_inline_markup(p, clean_inline_url(text))
            i += 1
            continue

        line = clean_inline_url(line)
        if line.startswith("دیده‌بان جهان نباید به عنوان"):
            add_callout(doc, line)
        else:
            p = doc.add_paragraph()
            set_paragraph_rtl(p)
            add_inline_markup(p, line)
        if line.startswith("جریان پیشنهادی:"):
            add_image(doc, IMG_SOURCES, "نمای مدیریت منابع در V15؛ این سطح باید از تجربه روزمره تحلیلگر جدا و برای نقش مدیر داده بهینه شود.", "اسکرین‌شات پنجره انتخاب و مدیریت منابع کرول در نسخه V15")
        i += 1


def audit_basic(doc):
    assert doc.sections[0].page_width == Inches(8.5)
    assert doc.sections[0].left_margin == Inches(1)
    assert doc.styles["Normal"].font.name == FONT
    for table in doc.tables:
        assert table._tbl.tblGrid is not None
        for row in table.rows:
            for cell in row.cells:
                assert cell._tc.tcPr.find(qn("w:tcW")) is not None


def main():
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_cover(doc)
    parse_body(doc, markdown)
    doc.core_properties.title = "دیده‌بان جهان — گزارش انبوردینگ محصول و تحلیل بنچمارک"
    doc.core_properties.subject = "معماری اطلاعات، تحلیل V15 و بنچمارک‌های World Monitor"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "دیده‌بان جهان، داشبورد اطلاعاتی، OSINT، معماری اطلاعات، AI"
    audit_basic(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
